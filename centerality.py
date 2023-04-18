from collections import defaultdict, deque
from queue import PriorityQueue
from docx import Document

from graph import Graph

class Centerality:

    def degree(self, graph):
        arr = []
        for node, neigbours in graph.adj_list.items():
            temp = 0

            for neighbour, cost in neigbours:
                temp += 1/cost

            arr.append((temp, node))

        arr.sort()

        dic = {}
        for cost, node in arr:
            dic[node] = cost

        return dic
        
    def betweenness(self, graph):
         # Initialize the betweenness centrality for all nodes to 0
        betweenness = {node: 0 for node in graph}

        for node in graph:
            # Initialize the distance and number of shortest paths for all nodes
            dist = {v: -1 for v in graph}
            num_paths = {v: 0 for v in graph}
            dist[node] = 0
            num_paths[node] = 1

            # Initialize a stack and a queue for depth-first search
            stack = [node]
            queue = []

            # Perform depth-first search and accumulate the number of shortest paths
            while stack:
                v = stack.pop()
                queue.append(v)
                for w, weight in graph[v]:
                    if dist[w] == -1:
                        dist[w] = dist[v] + weight
                        stack.append(w)
                    if dist[w] == dist[v] + weight:
                        num_paths[w] += num_paths[v]

            # Initialize the dependency and centrality for all nodes to 0
            dependency = {v: 0 for v in graph}
            centrality = {v: 0 for v in graph}

            # Traverse the queue in reverse order and accumulate the dependency and centrality
            while queue:
                w = queue.pop()
                for v, weight in graph[w]:
                    if dist[v] == dist[w] - weight:
                        dependency[v] += (num_paths[v] / num_paths[w]) * (1 + dependency[w])
                if w != node:
                    betweenness[w] += dependency[w]
                    centrality[w] += dependency[w]

                # Propagate the dependency and centrality to the predecessors of w
                for v, weight in graph[w]:
                    if dist[v] == dist[w] - weight:
                        dependency_v = (num_paths[v] / num_paths[w]) * (1 + dependency[w])
                        centrality[v] += dependency_v + centrality[w]
                        dependency[v] += dependency_v

        # Normalize the betweenness centrality by dividing by (n-1)*(n-2)/2
        n = len(graph)
        norm = (n-1) * (n-2) / 2
        betweenness = {node: bc / norm for node, bc in betweenness.items()}
     
        return betweenness


    def eigenvector(self, G, max_iter=300):
        prev_vals = dict.fromkeys(G.adj_list, 1)
        cur_vals = prev_vals.copy()
        
        for _ in range(max_iter):
            for node in G.adj_list:
                val = 0
                for neighbor in G.adj_list[node]:
                    val += prev_vals[neighbor[0]]
                cur_vals[node] = val
                
            # normalization
            norm = sum([i*i for i in cur_vals.values()])**0.5
            for val in cur_vals:
                cur_vals[val] /= norm
            prev_vals = cur_vals.copy()
    
        return cur_vals
    
    def pagerank(self, G, max_iter=100, d=0.85):
        # N = number of nodes
        N = len(G.adj_list)

        prev_rank = dict.fromkeys(G.adj_list, 1/N)
        cur_rank = prev_rank.copy()

        for _ in range(max_iter):
            for node in G.adj_list:
                val = 0
                for neighbor in G.adj_list[node]:
                    val += prev_rank[neighbor[0]] / len(G.adj_list[neighbor[0]])
                cur_rank[node] = (1-d) + d*val
            prev_rank = cur_rank.copy()
        return cur_rank

    def closenessCenterality(self, G, u=None, weighted=True, normalized=True):
        if weighted:
            path_length = UCS_tpl
        else:
            path_length = BFS_tpl

        if u is None:
            nodes = list(G.adj_list.keys())
        else:
            nodes = [u]

        closeness_centrality = {}
        for n in nodes:
            sp = path_length(G, n)  # sp = shortest path
            closeness_centrality[n] = 1/ sp
            if normalized:
                closeness_centrality[n] *= len(G.adj_list)-1

        if u is not None:
            return closeness_centrality[n]
        else:
            return closeness_centrality
        
    def katz(self, graph):
        alpha = 2
        beta = 1

        # Get the nodes and number of nodes in the graph
        nodes = sorted(graph.adj_list.keys())
        num_nodes = len(nodes)

        # Initialize the adjacency matrix
        adj_mat = [[0 for _ in range(num_nodes)] for _ in range(num_nodes)]

        # Fill the adjacency matrix with the weights of the edges
        for i, node in enumerate(nodes):
            for neighbor, weight in graph.adj_list[node]:
                j = nodes.index(neighbor)
                adj_mat[i][j] = weight
                adj_mat[j][i] = weight

        # Initialize the identity matrix
        I = [[1 if i == j else 0 for j in range(num_nodes)] for i in range(num_nodes)]

        # Calculate I - alpha*A
        I_minus_alpha_A = [[I[i][j] - alpha * adj_mat[i][j] for j in range(num_nodes)] for i in range(num_nodes)]

        # Initialize the inverse of I - alpha*A to zero
        inv_I_minus_alpha_A = [[0] * num_nodes for _ in range(num_nodes)]

        # Calculate the inverse of I - alpha*A using Gaussian elimination
        for i in range(num_nodes):
            # Initialize the row of the inverse matrix to zero except for the diagonal element
            inv_i = [0] * num_nodes
            inv_i[i] = 1

            # Perform Gaussian elimination to calculate the row of the inverse matrix
            for j in range(num_nodes):
                if j != i:
                    factor = I_minus_alpha_A[j][i] / I_minus_alpha_A[i][i]
                    for k in range(num_nodes):
                        I_minus_alpha_A[j][k] -= factor * I_minus_alpha_A[i][k]
                        inv_i[k] -= factor * inv_I_minus_alpha_A[i][k]

            # Normalize the row of the inverse matrix
            factor = 1 / I_minus_alpha_A[i][i]
            for j in range(num_nodes):
                inv_I_minus_alpha_A[i][j] = factor * inv_i[j]

        # Calculate the Katz centrality for each node
        katz = [[beta * inv_I_minus_alpha_A[i][j]] for i in range(num_nodes)]

        # Print the Katz centrality for each node
        dic = {}
        for i, node in enumerate(nodes):
            dic[node] = katz[i][0]

        print(dic)
        return dic


# UCS total path length from the start node to all of the nodes in the graph
def UCS_tpl(G, start):
    total_length = 0
    for end in G.adj_list:
        total_length += UCS_spl(G, start, end)
    return total_length

# ucs single path length calculator
def UCS_spl(G, start, goal):
    queue = PriorityQueue()
    queue.put((0, start, []))
    visted = set()
    while not queue.empty():
        cost, node, path = queue.get()
        if goal == node:
            return cost
        visted.add(node)
        for neighbour, edge_cost in G.adj_list[node]:
            if neighbour in visted:
                continue
            total_cost = cost + edge_cost
            queue.put((total_cost, neighbour, path + [node]))

    
# BFS total path length from the start node to all of the nodes in the graph
def BFS_tpl(G, start):
    total_length = 0
    for end in G.adj_list:
        total_length += BFS_spl(G, start, end)
    return total_length

# BFS single path length
def BFS_spl(G, start, goal):
        visted = {start}
        queue  = deque([(start, [start])])
        while queue:
            curr, path = queue.popleft()
            if curr == goal:
                return len(path)
            for neighbour, _ in G.adj_list[curr]:
                if neighbour not in visted:
                    visted.add(neighbour)
                    queue.append((neighbour, path +  [neighbour]))

if __name__ == "__main__":
    graph = Graph()
   
    Arad = graph.createNode('Arad', [('Sibiu', 140), ('Timisoara', 118), ('Zerind',75 )])
    Sibu = graph.createNode('Sibiu', [('Arad', 140), ('Oradea', 151), ('Fagaras', 99), ('Rimnicu Vilcea', 80)])
    Zerind = graph.createNode('Zerind', [('Oradea', 71), ('Arad',75 )])
    Timisora = graph.createNode('Timisoara', [('Arad', 118), ('Lugoj', 111)])
    Lugoj =  graph.createNode('Lugoj', [('Timisoara', 111), ('Mehadia', 70)])
    Mehadia = graph.createNode('Mehadia', [('Lugoj', 70), ('Drobeta', 75)])
    Drobeta = graph.createNode('Drobeta', [('Mehadia', 75), ('Craiova', 120)])
    Craiova = graph.createNode('Craiova', [('Drobeta', 120), ('Rimnicu Vilcea', 146), ('Pitesti', 138)])
    Rimnicu = graph.createNode('Rimnicu Vilcea', [('Sibiu', 80), ('Craiova', 146), ('Pitesti', 97)])
    Oradea = graph.createNode('Oradea', [('Sibiu', 151),('Zerind', 71) ])
    Fagaras = graph.createNode('Fagaras', [('Sibiu', 99), ('Bucharest', 211)])
    Pitesti = graph.createNode('Pitesti', [('Rimnicu Vilcea', 97), ('Craiova', 138), ('Bucharest', 101)])
    Bucharest = graph.createNode('Bucharest', [('Fagaras', 211), ('Pitesti', 101), ('Urziceni', 85), ('Giurgiu', 90)])
    Urziceni = graph.createNode('Urziceni', [('Vaslui', 142), ('Hirsova', 98), ('Bucharest', 85)])
    Hirsova = graph.createNode('Hirsova', [('Eforie', 86), ('Urziceni', 98)])
    Eforie = graph.createNode('Eforie', [('Hirsova', 86)])
    Vaslui = graph.createNode('Vaslui', [('Iasi', 92), ('Urziceni', 142)])
    Iasi = graph.createNode('Iasi', [('Vaslui', 92), ('Neamt', 87)])
    Neamt = graph.createNode('Neamt', [('Iasi', 87)])
    Giurgiu = graph.createNode('Giurgiu', [('Bucharest', 90)])
    
    nodes = [Eforie, Neamt, Iasi ,Giurgiu, Arad, Sibu, Zerind, Timisora, Lugoj, Mehadia, Drobeta, Craiova, Rimnicu, Oradea, Fagaras, Pitesti, Bucharest, Urziceni, Hirsova, Vaslui, Iasi]

    for node in nodes:
        graph.addNode(node)    
    
    # Open document to record values
    document = Document()
    
    mainTable = document.add_table(rows=1, cols=7)
    hdr_cells = mainTable.rows[0].cells
    
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Closeness'
    hdr_cells[2].text = 'Eigenvector'
    hdr_cells[3].text = 'Pagerank'
    hdr_cells[4].text = 'Degree'
    hdr_cells[5].text = 'Betweenness'
    hdr_cells[6].text = 'Katz'
    
    c = Centerality()
    ms = [c.closenessCenterality(graph), c.eigenvector(graph), c.pagerank(graph), c.degree(graph), c.betweenness(graph.adj_list), c.katz(graph)]
    ms_vals = [list(m.values()) for m in ms]
    
    nodes = list(graph.adj_list.keys())
    for i in range(len(nodes)):
        row_cells = mainTable.add_row().cells
        row_cells[0].text = nodes[i]
        for j in range(len(ms)):
            row_cells[j+1].text = str(ms_vals[j][i])
    
    document.add_page_break()
    
    document.add_heading('Top 5 cities', level=1)
    rankTable = document.add_table(rows=1, cols=6)
    hdr_cells = rankTable.rows[0].cells
    
    hdr_cells[0].text = 'Closeness'
    hdr_cells[1].text = 'Eigenvector'
    hdr_cells[2].text = 'Pagerank'
    hdr_cells[3].text = 'Degree'
    hdr_cells[4].text = 'Betweenness'
    hdr_cells[5].text = 'Katz'
    
    ranked = []
    for m in ms:
        ranked.append(sorted(m.items(), key= lambda item: item[1], reverse= True))
    
    top5s = []
    for tops in ranked:
        top5 = tops[:5]
        top5s.append([top[0] for top in top5])
    
    for i in range(5):
        row_cells = rankTable.add_row().cells
        for j in range(len(top5s)):
            row_cells[j].text = top5s[j][i]
    
    document.save('centerality.docx')