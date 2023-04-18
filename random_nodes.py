from docx import Document
from graph import Graph
from random import random, randint
from timeit import repeat

import matplotlib.pyplot as plt

ns = [10, 20, 30, 40]
ps = [0.2, 0.4, 0.6, 0.8]

loc = {}
for n in range(1, 41):
    while True:
        i, j = randint(0, 10), randint(0, 10)
        if (i, j) not in loc.values():
            loc[n] = (i, j)
            break

graphs = []
for n in ns:
    for p in ps:
        graph = Graph()
        graph.location = loc
        graph.addNodes([graph.createNode(val) for val in list(range(1, n+1))])
        for i in range(1, n+1):
            for j in range(i+1, n+1):
                prob = random()
                if prob < p:
                    graph.addNeighbour(i, j)
        
        graphs.append(graph)

# randomly select 10 nodes
random_nodes = []
for graph in graphs:
    rand_nodes = []
    if len(graph.adj_list) > 10:
        nodes = list(graph.adj_list.keys())
        for i in range(10):
            while True:
                ri = randint(0, len(graph.adj_list)-1)
                if nodes[ri] not in rand_nodes:
                    rand_nodes.append(nodes[ri])
                    break
    else:
        rand_nodes += list(graph.adj_list.keys())
    random_nodes.append(rand_nodes)


set_up = '''import __main__
from graph import Graph
from random import random, randint'''

bfs_code = '''
def test():
    graph.BFS(random_nodes[n][i], random_nodes[n][j])
'''

dfs_code = '''
def test():
    graph.DFS(random_nodes[n][i], random_nodes[n][j])
'''

astar_code = '''
def test():
    graph.astar(random_nodes[n][i], random_nodes[n][j])
'''

ucs_code = '''
def test():
    graph.UCS(random_nodes[n][i], random_nodes[n][j])
'''

greedy_code = '''
def test():
    graph[n].greedySearch(random_nodes[n][i], random_nodes[n][j])
'''

iter_code = '''
def test():
    graph[n].iterativeDeepening(random_nodes[n][i], random_nodes[n][j])
'''

codes = [dfs_code, ucs_code, astar_code, iter_code, bfs_code, greedy_code]

all_times = []
# find the path between them
for n in range(len(graphs)):
    g_time = []
    for code in codes:
        t = 0
        for i in range(10):
            for j in range(i+1, 10):
                t_av = sum(repeat(setup= set_up,
                       stmt= code,
                       repeat=5,
                       number= 1))/5
                t += t_av
        g_time.append(t)
    all_times.append(g_time)
# print(all_times)

all_paths = []
a_path = []
greed_path = []
dfs_path = []
iter_path = []

for n in range(len(graphs)):
    # for bfs, A* and UCS
    p = 0
    for i in range(10):
        for j in range(i+1, 10):
            p += len(graphs[n].astar(random_nodes[n][i], random_nodes[n][j]))-1
    a_path.append(p)
    # for greedy
    p = 0
    for i in range(10):
        for j in range(i+1, 10):
            p += len(graphs[n].greedySearch(random_nodes[n][i], random_nodes[n][j]))-1
    greed_path.append(p)
    # for iterativeDeepening
    p = 0
    for i in range(10):
        for j in range(i+1, 10):
            p += len(graphs[n].iterativeDeepening(random_nodes[n][i], random_nodes[n][j], 5))-1
    iter_path.append(p)
    # for dfs
    p = 0
    for i in range(10):
        for j in range(i+1, 10):
            p += len(graphs[n].DFS(random_nodes[n][i], random_nodes[n][j]))-1
    dfs_path.append(p)
all_paths.append(dfs_path)
all_paths.append(a_path)
all_paths.append(a_path)
all_paths.append(iter_path)
all_paths.append(a_path)
all_paths.append(greed_path)
# print(len(all_paths))
# print(all_times)



# create word doc and record the time and path length
document = Document()

# list all graphs
document.add_paragraph('List of all graphs').bold = True
for i in range(len(graphs)):
    p = document.add_paragraph(f'Graph {i+1}: {graphs[i].adj_list}')
    
document.add_paragraph('List of the selected random nodes').bold = True
for i in range(len(graphs)):
    p = document.add_paragraph(f'From Graph {i+1}: {random_nodes[i]}')

table = document.add_table(rows=1, cols=7)
hdr_cells = table.rows[0].cells

alg_names = ['DFS', 'UCS', 'A*', 'Iterative', 'BFS', 'Greedy']

hdr_cells[0].text = ''
for i in range(len(alg_names)):
    hdr_cells[i+1].text = alg_names[i]

for i in range(16):
    row_cells = table.add_row().cells
    row_cells[0].text = 'Graph '+str(i+1)
    for j in range(6):
        row_cells[j+1].text = str(all_times[i][j])

document.save('random_graph.docx')



# Graph ploting for time
arr = all_times

# define a list of colors for the circles
colors = ['red', 'blue', 'green', 'orange', 'purple', 'brown']

# define a dictionary to map algorithm index to label
algorithm_labels = {0: 'DFS', 1: 'UCS', 2: 'A*',
                    3: 'Iterative', 4: 'BFS', 5: 'Greedy'}

# plot circles for each algorithm and graph combination
fig, ax = plt.subplots()
for i in range(len(arr[0])):
    for j in range(len(arr)):
        x = j + 1
        y = arr[j][i]
        color = colors[i]
        if j == 0:
            label = algorithm_labels[i]
            ax.scatter(x, y, s=100, color=color, alpha=0.5, label=label)
        else:
            ax.scatter(x, y, s=100, color=color, alpha=0.5)

# set axis labels and legend
ax.set_xlabel('Number of Graphs')
ax.set_ylabel('Time')
ax.legend(title='Algorithms')
xticks = [j + 1 for j in range(len(arr))]
ax.set_xticks(xticks)
ax.set_xticklabels([f'{j+1}' for j in range(len(arr))])

# show the plot
plt.show()


# Graph ploting for paths
arr = all_paths

# define a list of colors for the circles
colors = ['red', 'blue', 'green', 'orange', 'purple', 'brown']

# define a dictionary to map algorithm index to label
algorithm_labels = {0: 'BFS', 1: 'UCS', 2: 'A*',
                    3: 'Greedy', 4: 'DFS', 5: 'Iterative'}

# plot circles for each algorithm and graph combination
fig, ax = plt.subplots()
for i in range(len(arr)):
    for j in range(len(arr[0])):
        x = j + 1
        y = arr[i][j]
        color = colors[i]
        if j == 0:
            label = algorithm_labels[i]
            ax.scatter(x, y, s=100, color=color, alpha=0.5, label=label)
        else:
            ax.scatter(x, y, s=100, color=color, alpha=0.5)

# set axis labels and legend
ax.set_xlabel('Number of Graphs')
ax.set_ylabel('Path Length')
ax.legend(title='Algorithms')
xticks = [j + 1 for j in range(len(arr[0]))]
ax.set_xticks(xticks)
ax.set_xticklabels([f'{j+1}' for j in range(len(arr[0]))])

# show the plot
plt.show()

print(all_paths)

# create word doc and record the path length
document = Document()

table = document.add_table(rows=1, cols=7)
hdr_cells = table.rows[0].cells

alg_names = ['DFS', 'UCS', 'A*', 'Iterative', 'BFS', 'Greedy']

hdr_cells[0].text = ''
for i in range(len(alg_names)):
    hdr_cells[i+1].text = alg_names[i]

for i in range(16):
    row_cells = table.add_row().cells
    row_cells[0].text = 'Graph '+str(i+1)
    for j in range(6):
        row_cells[j+1].text = str(all_paths[j][i])

document.save('random_graph_paths.docx')